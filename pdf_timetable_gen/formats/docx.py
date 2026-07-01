"""DOCX output writer."""

from __future__ import annotations

import logging
from datetime import date
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


def write(schedule, path: str | Path) -> Path:
    """Write schedule as a .docx file.

    Args:
        schedule: GeneratedSchedule.
        path: Output file path.

    Returns:
        Path to the written file.
    """
    path = Path(path)
    doc = Document()

    # --- Title ---
    title = doc.add_heading("Study Timetable", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Meta ---
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {date.today().strftime('%Y-%m-%d')} | "
                 f"Days: {len(schedule.days)} | "
                 f"Topics: {schedule.total_topics} | "
                 f"Hours: {schedule.total_hours:.1f}")

    doc.add_paragraph("")  # spacer

    # --- Daily plans ---
    for day in schedule.days:
        date_str = day.date.strftime("%Y-%m-%d (%a)") if day.date else f"Day {day.day_number}"
        doc.add_heading(f"Day {day.day_number} — {date_str}", level=1)

        # Table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        hdr[0].text = "Phase"
        hdr[1].text = "Topic"
        hdr[2].text = "Chapter"
        hdr[3].text = "Hours"

        phase_icons = {"learn": "📖", "practice": "✏️", "revise": "🔄"}
        for slot in day.slots:
            row = table.add_row().cells
            row[0].text = f"{phase_icons.get(slot.phase, '📚')} {slot.phase}"
            row[1].text = slot.topic_title
            row[2].text = slot.chapter
            row[3].text = f"{slot.hours:.1f}h"

        doc.add_paragraph(f"Total: {day.total_hours:.1f} hours")
        doc.add_paragraph("")  # spacer between days

    doc.save(str(path))
    logger.info("Wrote DOCX: %s", path)
    return path
